from __future__ import annotations

# -------------------------
# Built-in
# -------------------------
from dataclasses import dataclass
from decimal import Decimal
from pathlib import Path
from typing import Any, Dict, List, Optional, Tuple
from datetime import datetime

# -------------------------
# Terceiros
# -------------------------
from sqlalchemy.orm import Session

from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn

from app.db.models import EfdRevisao
# -------------------------
# Projeto (utils)
# -------------------------
from app.sped.blocoM.m_utils import (
    _clean_sped_line,
    _reg_of_line,
    _d,
    _fmt_br,
)

# ============================================================
# Layout / Parser (layout-driven)
# (Você já tem isso no projeto; repliquei aqui pra ficar completo)
# ============================================================

def _parse_linha_sped_to_reg_dados(linha: str) -> Tuple[str, List[Any]]:
    s = (linha or "").strip()
    if not s:
        raise ValueError("Linha SPED vazia")
    s = s.lstrip("\ufeff")
    anchor = "|C170|"
    p = s.find(anchor)
    if p >= 0:
        s = s[p:]
    if not s.startswith("|"):
        p2 = s.find("|")
        if p2 >= 0:
            s = s[p2:]
    s = s.strip()
    if not (s.startswith("|") and "|" in s[1:]):
        raise ValueError(f"Linha SPED inválida: {s[:80]}")
    parts = s.strip().strip("|").split("|")
    reg = parts[0].strip().upper()
    dados = parts[1:]
    return reg, dados


class C170Layout:
    idx_cfop: int = 9
    idx_cst_pis: int = 23
    idx_cst_cofins: int = 29

    idx_vl_item: int = 5
    idx_vl_bc_pis: int = 24
    idx_aliq_pis: int = 25
    idx_vl_pis: int = 28

    idx_vl_bc_cofins: int = 30
    idx_aliq_cofins: int = 31
    idx_vl_cofins: int = 34


LAYOUT_C170 = C170Layout()

# -------------------------
# DTO
# -------------------------
@dataclass(frozen=True)
class DossieExportacaoDados:
    # Identificação
    empresa_nome: str
    cnpj: str
    periodo: str              # MM/AAAA
    periodo_yyyymm: str       # YYYYMM

    # Exportação (Bloco C) — layout-driven (C170)
    export_base_total: Decimal
    export_itens: int
    export_por_cfop: Dict[str, Decimal]

    # ✅ Ajuste (AJUSTE_M) — rastreio (DB)
    ajuste_tipo: Optional[str] = None
    ajuste_base_exportacao: Decimal = Decimal("0.00")
    ajuste_credito_extra_pis: Decimal = Decimal("0.00")
    ajuste_credito_extra_cofins: Decimal = Decimal("0.00")
    ajuste_credito_extra_total: Decimal = Decimal("0.00")
    ajuste_origem: Optional[str] = None          # ex.: EXP_RESSARC_V1
    ajuste_created_at: Optional[str] = None      # string formatada

    # Apuração (Bloco M) — pode não existir
    tem_m100: bool = False
    tem_m500: bool = False
    base_credito_mes: Decimal = Decimal("0.00")
    credito_pis: Decimal = Decimal("0.00")
    credito_cofins: Decimal = Decimal("0.00")
    credito_total: Decimal = Decimal("0.00")
    aliq_pis_pct: str = "—"
    aliq_cofins_pct: str = "—"

    # Composição por CST (M105/M505)
    csts: List[str] = None
    base_por_cst: Dict[str, Decimal] = None

    # Evidências e observações
    bloco_m_linhas_evidencia: List[str] = None
    observacoes: List[str] = None

    # ✅ Evidências bloco 1 e 0900
    linha_0900: Optional[str] = None
    bloco_1_linhas_evidencia: List[str] = None

    # Fonte do TXT (para o DOCX)
    fonte_txt_path: Path = None
# -------------------------
# Helpers
# -------------------------

def _safe_str(x: Any) -> str:
    return (str(x) if x is not None else "").strip()


def _fmt_pct_from_br(v: str) -> str:
    s = (v or "").strip()
    if not s:
        return ""
    try:
        d = _d(s)
        return f"{d.quantize(Decimal('0.01'))}%".replace(".", ",")
    except Exception:
        return s + "%"


def _extrair_fields(line: str) -> Tuple[str, List[str]]:
    ln = _clean_sped_line(line or "")
    parts = ln.strip("|").split("|")
    if not parts:
        return "", []
    return parts[0].upper(), parts[1:]


def _find_first(lines: List[str], reg: str) -> Optional[str]:
    for l in lines or []:
        if _reg_of_line(l) == reg:
            return _clean_sped_line(l)
    return None


def _find_all(lines: List[str], reg: str) -> List[str]:
    return [_clean_sped_line(l) for l in (lines or []) if _reg_of_line(l) == reg]


def _parse_m100_m500_credit_base_aliq(line: str) -> Tuple[Decimal, Decimal, str]:
    """
    Retorna: (base_vl_bc, credito_vl_cred, aliq_pct_str)
    No layout usado por vocês:
      - VL_BC = dados[2]
      - ALIQ  = dados[3]
      - VL_CRED = dados[6]
    """
    _, dados = _extrair_fields(line)
    base = Decimal("0")
    cred = Decimal("0")
    aliq = ""

    try:
        if len(dados) >= 3:
            base = _d(dados[2] or "0")
        if len(dados) >= 4:
            aliq = _fmt_pct_from_br(dados[3] or "")
        if len(dados) >= 7:
            cred = _d(dados[6] or "0")
    except Exception:
        pass

    return base, cred, aliq


def _parse_m105_m505(line: str) -> Tuple[Optional[str], Decimal]:
    """
    |M105|NAT_BC_CRED|CST|VL_BC|...|
    |M505|NAT_BC_CRED|CST|VL_BC|...|
    """
    _, dados = _extrair_fields(line)
    cst = None
    base = Decimal("0")
    try:
        if len(dados) >= 3:
            cst = _safe_str(dados[1])      # ✅ CST
            base = _d(dados[2] or "0")     # ✅ VL_BC
    except Exception:
        pass
    return cst, base


def _ler_txt_lines(p: Path) -> List[str]:
    try:
        return p.read_text(encoding="utf-8").splitlines()
    except Exception:
        return p.read_text(encoding="latin-1").splitlines()


def _extrair_periodo_0000(linhas: List[str]) -> Tuple[str, str]:
    """
    Retorna (MM/AAAA, YYYYMM) a partir de DT_INI do 0000 (DDMMAAAA).
    """
    for ln in linhas or []:
        s = (ln or "").strip()
        if not s.startswith("|0000|"):
            continue
        parts = s.strip("|").split("|")
        dt_ini = parts[5] if len(parts) > 5 else ""
        dig = "".join(c for c in (dt_ini or "") if c.isdigit())
        if len(dig) == 8:
            mm = dig[2:4]
            yyyy = dig[4:8]
            return f"{mm}/{yyyy}", f"{yyyy}{mm}"
        break
    return "", ""


def _extrair_cnpj_0140_0100(linhas: List[str]) -> str:
    """
    Procura um token com 14 dígitos no 0140 ou 0100.
    """
    for prefix in ("|0140|", "|0100|"):
        for ln in linhas or []:
            s = (ln or "").strip()
            if not s.startswith(prefix):
                continue
            parts = s.strip("|").split("|")
            for tok in parts:
                dig = "".join(c for c in (tok or "") if c.isdigit())
                if len(dig) == 14:
                    return dig
    return ""


def _extrair_linhas_m(linhas: List[str]) -> List[str]:
    return [_clean_sped_line(ln) for ln in (linhas or []) if (ln or "").lstrip().startswith("|M")]


def _is_cfop_export(cfop: str) -> bool:
    c = "".join(ch for ch in (cfop or "") if ch.isdigit())
    return len(c) == 4 and c.startswith("7")


def _extrair_exportacao_c170_layout(linhas: List[str]) -> Tuple[Decimal, Dict[str, Decimal], int]:
    """
    Exportação layout-driven via C170:
      - CFOP = LAYOUT_C170.idx_cfop
      - Base = VL_ITEM = LAYOUT_C170.idx_vl_item
    Critério: CFOP 7xxx.
    """
    total = Decimal("0")
    por_cfop: Dict[str, Decimal] = {}
    itens = 0

    for ln in linhas or []:
        if "|C170|" not in (ln or ""):
            continue

        try:
            reg, dados = _parse_linha_sped_to_reg_dados(ln)
        except Exception:
            continue
        if reg != "C170":
            continue

        # cfop
        if len(dados) <= LAYOUT_C170.idx_cfop:
            continue
        cfop = (dados[LAYOUT_C170.idx_cfop] or "").strip()
        if not _is_cfop_export(cfop):
            continue

        # vl_item
        if len(dados) <= LAYOUT_C170.idx_vl_item:
            continue
        try:
            vl_item = _d(dados[LAYOUT_C170.idx_vl_item] or "0")
        except Exception:
            continue
        if vl_item <= 0:
            continue

        total += vl_item
        por_cfop[cfop] = por_cfop.get(cfop, Decimal("0")) + vl_item
        itens += 1

    return (
        total.quantize(Decimal("0.01")),
        {k: v.quantize(Decimal("0.01")) for k, v in por_cfop.items()},
        itens,
    )
def _extrair_linha_0900(linhas: List[str]) -> Optional[str]:
    ln = _find_first(linhas, "0900")
    return _clean_sped_line(ln) if ln else None

def _extrair_linhas_bloco_1_evidencia(linhas: List[str]) -> List[str]:
    regs = {"1100", "1500", "1990"}
    out = []
    for ln in linhas or []:
        r = _reg_of_line(ln)
        if r in regs:
            out.append(_clean_sped_line(ln))
    return out


def _pick_txt_from_dir(dir_path: Path, prefer_name_contains: Optional[str] = None) -> Path:
    if not dir_path.exists():
        raise ValueError(f"Pasta não encontrada: {dir_path}")

    arquivos = [p for p in dir_path.glob("*.txt") if p.is_file()]
    if not arquivos:
        raise ValueError(f"Nenhum .txt encontrado em: {dir_path}")

    # 0) Preferência explícita
    if prefer_name_contains:
        prefer = prefer_name_contains.upper()
        preferidos = [p for p in arquivos if prefer in p.stem.upper()]
        if preferidos:
            return sorted(preferidos, key=lambda p: p.stat().st_mtime, reverse=True)[0]

    # 1) Prioridade padrão: RETIFICADO
    retificados = [p for p in arquivos if "_RETIFICADO" in p.stem.upper()]
    if retificados:
        return sorted(retificados, key=lambda p: p.stat().st_mtime, reverse=True)[0]

    # 2) Senão, mais recente
    return sorted(arquivos, key=lambda p: p.stat().st_mtime, reverse=True)[0]

def _buscar_ajuste_exportacao_db(db: Session, versao_id: int) -> Optional[Dict[str, Any]]:
    """
    Busca o AJUSTE_M mais recente ligado à versão (preferência: versao_revisada_id = versao_id).
    """
    q = (
        db.query(EfdRevisao)
        .filter(EfdRevisao.acao == "AJUSTE_M")
        .filter(
            (EfdRevisao.versao_revisada_id == int(versao_id)) |
            (EfdRevisao.versao_origem_id == int(versao_id))
        )
        .order_by(EfdRevisao.id.desc())
    )
    rev = q.first()
    if not rev:
        return None

    rj = dict(rev.revisao_json or {})
    meta = rj.get("meta") if isinstance(rj.get("meta"), dict) else {}
    detalhe = rj.get("detalhe") if isinstance(rj.get("detalhe"), dict) else {}

    return {
        "meta": meta,
        "detalhe": detalhe,
        "motivo_codigo": rev.motivo_codigo,
        "created_at": rev.created_at,
        "apontamento_id": rev.apontamento_id,
        "id": rev.id,
    }

# -------------------------
# Core: TXT -> Dados
# -------------------------

def montar_dados_dossie_exportacao_de_txt(
    txt_path: Path,
    *,
    empresa_nome_override: Optional[str] = None,
    ajuste_db: Optional[Dict[str, Any]] = None,
) -> DossieExportacaoDados:
    linhas = _ler_txt_lines(txt_path)
    linha_0900 = _extrair_linha_0900(linhas)
    bloco_1_evid = _extrair_linhas_bloco_1_evidencia(linhas)

    periodo_mm_aaaa, periodo_yyyymm = _extrair_periodo_0000(linhas)
    cnpj = _extrair_cnpj_0140_0100(linhas) or "—"
    empresa_nome = empresa_nome_override or "—"

    # Exportação layout-driven (C170)
    export_total, export_por_cfop, export_itens = _extrair_exportacao_c170_layout(linhas)

    # ✅ AJUSTE_M (delta) — vem do DB (se fornecido)
    ajuste_tipo = None
    ajuste_base_exportacao = Decimal("0.00")
    ajuste_origem = None
    ajuste_created_at = None

    if isinstance(ajuste_db, dict):
        meta = ajuste_db.get("meta") or {}
        if isinstance(meta, dict):
            ajuste_tipo = _safe_str(meta.get("tipo")) or None
            ajuste_origem = _safe_str(meta.get("origem_regra")) or _safe_str(ajuste_db.get("motivo_codigo")) or None
            ajuste_base_exportacao = _d(meta.get("base_exportacao") or "0").quantize(Decimal("0.01"))

        ca = ajuste_db.get("created_at")
        if ca:
            try:
                ajuste_created_at = ca.strftime("%d/%m/%Y %H:%M:%S")
            except Exception:
                ajuste_created_at = str(ca)

    ajuste_extra_pis = (ajuste_base_exportacao * Decimal("0.0165")).quantize(Decimal("0.01"))
    ajuste_extra_cof = (ajuste_base_exportacao * Decimal("0.0760")).quantize(Decimal("0.01"))
    ajuste_extra_total = (ajuste_extra_pis + ajuste_extra_cof).quantize(Decimal("0.01"))

    # Bloco M (linhas)
    bloco_m_raw = _extrair_linhas_m(linhas)
    if bloco_m_raw:
        corpo = [ln for ln in bloco_m_raw if _reg_of_line(ln) not in ("M001", "M990")]
        bloco_m_final = ["|M001|0|"] + corpo + [f"|M990|{len(corpo) + 2}|"]
    else:
        bloco_m_final = ["|M001|0|", "|M990|2|"]

    m100 = _find_first(bloco_m_final, "M100")
    m500 = _find_first(bloco_m_final, "M500")
    m105s = _find_all(bloco_m_final, "M105")
    m505s = _find_all(bloco_m_final, "M505")

    tem_m100 = bool(m100)
    tem_m500 = bool(m500)

    # Observações
    observacoes: List[str] = [f"Fonte: TXT: {txt_path.name}"]

    if export_total > 0:
        cfops = ", ".join(sorted(export_por_cfop.keys())[:10])
        observacoes.append(f"✅ Exportação detectada (CFOP 7xxx) no Bloco C (C170): base R$ {_fmt_br(export_total)} | CFOPs: {cfops}")
    else:
        observacoes.append("⚠️ Não identifiquei CFOP 7xxx no C170 (Bloco C) no TXT.")

    if not tem_m100 and not tem_m500:
        observacoes.append("ℹ️ Sem M100/M500 no período: não há apuração de créditos no Bloco M (isso pode ser coerente).")

    # Parse do crédito (quando existir)
    base_pis, cred_pis, aliq_pis = _parse_m100_m500_credit_base_aliq(m100 or "")
    base_cof, cred_cof, aliq_cof = _parse_m100_m500_credit_base_aliq(m500 or "")

    credito_total = (cred_pis + cred_cof).quantize(Decimal("0.01"))
    base_credito_mes = (base_pis if base_pis > 0 else base_cof).quantize(Decimal("0.01"))

    # Composição por CST (preferencialmente M105; fallback M505)
    base_por_cst: Dict[str, Decimal] = {}
    for ln in (m105s or []):
        cst, b = _parse_m105_m505(ln)
        if cst and b > 0:
            base_por_cst[cst] = base_por_cst.get(cst, Decimal("0.00")) + b

    if not base_por_cst:
        for ln in (m505s or []):
            cst, b = _parse_m105_m505(ln)
            if cst and b > 0:
                base_por_cst[cst] = base_por_cst.get(cst, Decimal("0.00")) + b

    # Evidências (trechos M)
    evid: List[str] = []
    for reg in ("M100", "M105", "M500", "M505"):
        for ln in bloco_m_final or []:
            if _reg_of_line(ln) == reg:
                evid.append(_clean_sped_line(ln))

    csts = sorted(base_por_cst.keys())

    return DossieExportacaoDados(
        empresa_nome=empresa_nome,
        cnpj=cnpj,
        periodo=periodo_mm_aaaa or "—",
        periodo_yyyymm=periodo_yyyymm or "—",

        export_base_total=export_total,
        export_itens=int(export_itens),
        export_por_cfop=export_por_cfop,
        ajuste_tipo=ajuste_tipo,
        ajuste_base_exportacao=ajuste_base_exportacao,
        ajuste_credito_extra_pis=ajuste_extra_pis,
        ajuste_credito_extra_cofins=ajuste_extra_cof,
        ajuste_credito_extra_total=ajuste_extra_total,
        ajuste_origem=ajuste_origem,
        ajuste_created_at=ajuste_created_at,

        tem_m100=tem_m100,
        tem_m500=tem_m500,
        base_credito_mes=base_credito_mes,
        credito_pis=cred_pis.quantize(Decimal("0.01")),
        credito_cofins=cred_cof.quantize(Decimal("0.01")),
        credito_total=credito_total,
        aliq_pis_pct=aliq_pis or "—",
        aliq_cofins_pct=aliq_cof or "—",

        csts=csts,
        base_por_cst={k: v.quantize(Decimal("0.01")) for k, v in base_por_cst.items()},
        linha_0900=linha_0900,
        bloco_1_linhas_evidencia=bloco_1_evid,
        bloco_m_linhas_evidencia=evid,
        observacoes=observacoes,
        fonte_txt_path=txt_path,
    )


# -------------------------
# DOCX
# -------------------------

def gerar_docx_dossie_exportacao(*, dados: DossieExportacaoDados, output_dir: Optional[Path] = None) -> Path:
    if output_dir is None:
        output_dir = Path.home() / "Downloads" / "Dossies"
    output_dir.mkdir(parents=True, exist_ok=True)

    # Nome baseado no TXT em questão
    stem = dados.fonte_txt_path.stem.replace("_RETIFICADO", "")
    fname = f"dossie_{stem}_RETIFICADO_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
    path = output_dir / fname

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    style.font.size = Pt(11)

    doc.add_heading("Dossiê Técnico — Exportação e Créditos PIS/COFINS (SPED)", level=1)
    doc.add_paragraph(f"Empresa: {dados.empresa_nome}")
    doc.add_paragraph(f"CNPJ: {dados.cnpj}")
    doc.add_paragraph(f"Período: {dados.periodo} (YYYYMM={dados.periodo_yyyymm})")
    doc.add_paragraph("")

    # 1) Enquadramento coerente
    doc.add_heading("1) Enquadramento (resumo)", level=2)
    doc.add_paragraph(
        "A exportação (CFOP 7xxx) pode levar a acúmulo de créditos de PIS/COFINS oriundos das entradas/insumos. "
        "O ressarcimento/compensação depende da apuração e escrituração correta no Bloco M (M100/M105 para PIS; "
        "M500/M505 para COFINS) e dos controles do Bloco 1 (ex.: 1100/1500 e, quando aplicável, 1200/1210/1700). "
        "Este dossiê consolida evidências do TXT do SPED para apoiar a revisão fiscal e eventual PER/DCOMP Web."
    )

    # 2) Evidências de exportação
    doc.add_heading("2) Evidências de Exportação (Bloco C — C170)", level=2)
    if dados.export_base_total > 0:
        doc.add_paragraph(f"Base exportação (layout-driven via C170/CFOP 7xxx): R$ {_fmt_br(dados.export_base_total)}")
        if dados.export_por_cfop:
            table = doc.add_table(rows=1, cols=2)
            hdr = table.rows[0].cells
            hdr[0].text = "CFOP"
            hdr[1].text = "Base (R$)"
            for cfop in sorted(dados.export_por_cfop.keys())[:50]:
                row = table.add_row().cells
                row[0].text = cfop
                row[1].text = _fmt_br(dados.export_por_cfop[cfop])
        doc.add_paragraph(f"Itens C170 considerados: {dados.export_itens}")
    else:
        doc.add_paragraph("Não foi possível identificar exportação por CFOP 7xxx no C170 do TXT (Bloco C).")

    # 3) Apuração no Bloco M
    doc.add_heading("3) Apuração de Créditos (Bloco M)", level=2)
    if dados.tem_m100 or dados.tem_m500:
        doc.add_paragraph(f"Base (VL_BC — Bloco M): R$ {_fmt_br(dados.base_credito_mes)}")
        doc.add_paragraph(f"PIS: alíquota {dados.aliq_pis_pct} | crédito R$ {_fmt_br(dados.credito_pis)}")
        doc.add_paragraph(f"COFINS: alíquota {dados.aliq_cofins_pct} | crédito R$ {_fmt_br(dados.credito_cofins)}")
        doc.add_paragraph(f"TOTAL (PIS+COFINS): R$ {_fmt_br(dados.credito_total)}")
    else:
        doc.add_paragraph(
            "Não foi identificada apuração de créditos no período (sem M100/M500). "
            "Isso pode ser coerente quando não há crédito acumulado escriturado. "
            "Em caso de expectativa de ressarcimento por exportação, revisar as entradas/insumos e a apuração/escrituração do Bloco M."
        )

    doc.add_heading("3.1) Ajuste de Exportação (AJUSTE_M)", level=3)
    aj_base = getattr(dados, "ajuste_base_exportacao", Decimal("0.00"))
    if aj_base > 0:
        doc.add_paragraph(f"Tipo: {dados.ajuste_tipo or '—'} | Origem: {dados.ajuste_origem or '—'}")
        if dados.ajuste_created_at:
            doc.add_paragraph(f"Registrado em: {dados.ajuste_created_at}")
        doc.add_paragraph(f"Base exportação (delta): R$ {_fmt_br(dados.ajuste_base_exportacao)}")
        doc.add_paragraph(f"Crédito extra estimado: R$ {_fmt_br(dados.ajuste_credito_extra_total)} "
                          f"(PIS {_fmt_br(dados.ajuste_credito_extra_pis)} + COFINS {_fmt_br(dados.ajuste_credito_extra_cofins)})")
    else:
        doc.add_paragraph("Não foi identificado AJUSTE_M de exportação no banco para esta versão.")


    # 4) Composição por CST
    doc.add_heading("4) Composição por CST (M105/M505)", level=2)
    if dados.base_por_cst:
        table = doc.add_table(rows=1, cols=2)
        hdr = table.rows[0].cells
        hdr[0].text = "CST"
        hdr[1].text = "Base (R$)"
        for cst in dados.csts[:50]:
            row = table.add_row().cells
            row[0].text = str(cst)
            row[1].text = _fmt_br(dados.base_por_cst[cst])
    else:
        doc.add_paragraph("Não foi possível identificar bases por CST via M105/M505 no TXT.")

    # 5) Evidências no Bloco M
    doc.add_heading("5) Evidências no SPED (trechos do Bloco M)", level=2)
    if (dados.bloco_m_linhas_evidencia or []):
        if "EvidenciaMonospace" not in [s.name for s in doc.styles]:
            evid_style = doc.styles.add_style("EvidenciaMonospace", 1)
            evid_style.font.name = "Consolas"
            evid_style._element.rPr.rFonts.set(qn("w:eastAsia"), "Consolas")
            evid_style.font.size = Pt(9)

        for ln in dados.bloco_m_linhas_evidencia[:200]:
            doc.add_paragraph(_clean_sped_line(ln), style="EvidenciaMonospace")
    else:
        doc.add_paragraph("Sem linhas M100/M105/M500/M505 para evidência (Bloco M sem apuração de crédito no período).")

    doc.add_heading("5.1) Evidência de Receitas (0900)", level=3)
    if dados.linha_0900:
        doc.add_paragraph(dados.linha_0900)
        doc.add_paragraph("Observação: o ajuste aplicado afeta o Bloco M (crédito), não altera as receitas do período.")
    else:
        doc.add_paragraph("Linha 0900 não encontrada no TXT.")

    doc.add_heading("5.2) Evidências do Bloco 1 (1100/1500)", level=3)
    if (dados.bloco_1_linhas_evidencia or []):
        for ln in dados.bloco_1_linhas_evidencia[:200]:
            doc.add_paragraph(_clean_sped_line(ln),
                              style="EvidenciaMonospace" if "EvidenciaMonospace" in [s.name for s in
                                                                                     doc.styles] else None)
    else:
        doc.add_paragraph("Sem evidências do Bloco 1 (1100/1500) no TXT.")

    # 6) Observações
    if dados.observacoes:
        doc.add_heading("6) Observações", level=2)
        for obs in dados.observacoes:
            doc.add_paragraph(f"- {obs}")

    doc.add_paragraph(f"Gerado em {datetime.now().strftime('%d/%m/%Y %H:%M:%S')}")
    doc.add_heading("7) Checklist PER/DCOMP Web (operacional)", level=2)
    doc.add_paragraph("- Transmitir a EFD-Contribuições retificadora e aguardar processamento na Receita.")
    doc.add_paragraph("- No PER/DCOMP Web: selecionar crédito de PIS/COFINS conforme apuração do período (Bloco M).")
    doc.add_paragraph(
        "- Em caso de intimação/solicitação: anexar este dossiê + memória de cálculo + evidências de exportação (CFOP 7xxx).")
    doc.add_paragraph(
        "- Manter rastreabilidade do ajuste: motivo/regra (EXP_RESSARC_V1) e referência do SPED retificado.")
    doc.save(str(path))
    return path


# -------------------------
# API-friendly wrappers (TXT source of truth)
# -------------------------

def montar_dados_dossie_exportacao(
    db: Session,
    *,
    versao_id: int,
    empresa_nome_override: Optional[str] = None,
) -> DossieExportacaoDados:
    """
    Fonte da verdade: TXT mais recente na pasta Downloads/Dossies.
    versao_id é mantido só por compatibilidade.
    """
    pasta = Path.home() / "Downloads" / "Dossies"
    txt_path = _pick_txt_from_dir(dir_path=pasta)
    return montar_dados_dossie_exportacao_de_txt(txt_path, empresa_nome_override=empresa_nome_override)


def gerar_dossie_exportacao_docx(
    db: Session,
    *,
    versao_id: int,
    empresa_nome_override: Optional[str] = None,
    output_dir: Optional[Path] = None,
) -> Path:
    # ✅ mantém exatamente o comportamento antigo (TXT na pasta Dossies)
    pasta = Path.home() / "Downloads" / "Dossies"
    txt_path = _pick_txt_from_dir(dir_path=pasta)

    # ✅ incrementa com AJUSTE_M do banco (se existir)
    ajuste_db = _buscar_ajuste_exportacao_db(db, versao_id=int(versao_id))

    dados = montar_dados_dossie_exportacao_de_txt(
        txt_path,
        empresa_nome_override=empresa_nome_override,
        ajuste_db=ajuste_db,
    )
    return gerar_docx_dossie_exportacao(dados=dados, output_dir=output_dir)


def gerar_dossie_exportacao_docx_da_pasta(
    *,
    pasta_txt: Path,
    output_dir: Optional[Path] = None,
    empresa_nome_override: Optional[str] = None,
    prefer_name_contains: Optional[str] = None,
) -> Path:
    """
    Gera DOCX a partir de um TXT escolhido dentro de uma pasta específica.
    """
    txt_path = _pick_txt_from_dir(dir_path=pasta_txt, prefer_name_contains=prefer_name_contains)
    dados = montar_dados_dossie_exportacao_de_txt(txt_path, empresa_nome_override=empresa_nome_override)
    return gerar_docx_dossie_exportacao(dados=dados, output_dir=output_dir)
